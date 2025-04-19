# report_generator.py (Version avec injection LLM LangChain + Fonction d'assemblage)

import config
# Importer le type LangChain
from langchain_core.language_models.chat_models import BaseChatModel
from vector_database import VectorDBManager
from data_models import ReportPlan, ReportSection # Assurez-vous que ReportSection est importé
from memory_manager import MemoryManager # <--- AJOUT : Pour charger le plan
from typing import List, Optional
import logging
from docx import Document
from docx.shared import Inches, Pt # <--- AJOUT : Pt pour la taille de police
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # <--- AJOUT : Pour l'alignement

log = logging.getLogger(__name__)
if not log.handlers: logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - [%(module)s] - %(message)s')

class ReportGenerator:
    """Generates the report content section by section (Non-Agentic Command)."""

    def __init__(self, vector_db: VectorDBManager, llm_instance: BaseChatModel):
        """Initializes with VectorDB and LangChain LLM instances."""
        if not vector_db: raise ValueError("VectorDBManager instance required.")
        if not llm_instance: raise ValueError("LLM instance required.")
        self.vector_db = vector_db
        self.llm = llm_instance # Utilise l'instance injectée
        log.info("ReportGenerator initialized with VectorDB and LLM.")

    def _get_context_for_section(self, section_title: str, k_guidelines: int = 3, k_journals: int = 7) -> str:
        """Retrieves relevant text chunks from vector DBs for a section title."""
        log.debug(f"Getting context for section: '{section_title}'")
        guideline_context = "Not searched."; journal_context = "Not searched."
        try:
            guideline_results = self.vector_db.search_references(query=section_title, k=k_guidelines)
            if guideline_results: guideline_context = "\n---\n".join([r.get("document","") for r in guideline_results])
            else: guideline_context = "No relevant guidelines found."
        except Exception as e: guideline_context = f"Error searching guidelines: {e}"; log.error(e, exc_info=True)
        try:
            journal_results = self.vector_db.search_journals(query=section_title, k=k_journals)
            if journal_results: journal_context = "\n---\n".join([r.get("document","") for r in journal_results])
            else: journal_context = "No relevant journal entries found."
        except Exception as e: journal_context = f"Error searching journals: {e}"; log.error(e, exc_info=True)

        combined = f"GUIDELINES CONTEXT:\n{guideline_context}\n\nJOURNAL ENTRIES CONTEXT:\n{journal_context}"
        log.debug(f"Context length for '{section_title}': {len(combined)}")
        return combined


    def _generate_section_content(self, section: ReportSection, report_plan: ReportPlan) -> Optional[str]:
        """Generates content for a single section using LLM and context."""
        section_title = getattr(section, 'title', 'Untitled Section')
        log.info(f"Generating content for section: '{section_title}'...")

        # Gérer sections spéciales (Biblio, Appendices n'ont pas besoin de génération LLM ici)
        if section_title.lower() in ["bibliography", "appendices (optional)"]:
             log.info(f"Skipping LLM generation for section '{section_title}'.")
             # On pourrait mettre un placeholder ou retourner None pour que generate_full_report l'ignore
             return f"({section_title} content to be added manually or via ReferenceManager)."

        # 1. Récupérer le contexte
        context = self._get_context_for_section(section_title)

        # 2. Préparer le prompt pour Gemini (via LangChain)
        instructions = f"Draft the content for the report section titled '{section_title}'. Use the provided context. Maintain a professional and academic tone. Synthesize information."
        system_instructions = "You are an AI assistant writing a section for a professional MSc report based *only* on provided context."
        user_prompt = f"{instructions}\n\nCONTEXT:\n---\n{context[:15000]}\n---\n\nDraft for section \"{section_title}\":"
        full_prompt = f"System Instructions:\n{system_instructions}\n\nUser Request:\n{user_prompt}"

        # 3. Appeler le LLM via LangChain
        try:
            response_msg = self.llm.invoke(full_prompt)
            drafted_content = getattr(response_msg, 'content', None)

            if drafted_content:
                log.info(f"Successfully drafted content for '{section_title}'. Length: {len(drafted_content)}.")
                return drafted_content
            else:
                log.error(f"LLM returned empty or invalid response for section '{section_title}'. Response: {response_msg}")
                return None # Échec de la génération
        except Exception as e:
            log.error(f"LLM call failed during drafting section '{section_title}': {e}", exc_info=True)
            return None # Échec de la génération


    def generate_full_report(self, report_plan: ReportPlan, output_docx_path: str = config.DEFAULT_REPORT_OUTPUT) -> ReportPlan:
        """Generates content for all sections and saves to a DOCX file."""
        log.info("Starting full report generation (non-agentic)...")
        document = Document()
        # Utiliser le titre du plan s'il existe
        report_title = getattr(report_plan, 'title', "Apprenticeship Report")
        document.add_heading(report_title, level=0)

        # Fonction récursive pour traiter les sections
        def process_and_add_section(section: ReportSection):
            # Générer le contenu seulement si pas déjà présent ou si statut est 'pending'/'failed'?
            # Pour cette commande non-agentique, on regénère tout.
            content = self._generate_section_content(section, report_plan)
            section.content = content # Stocke le contenu (ou None) dans l'objet plan
            if content:
                 section.status = "drafted"
                 try:
                     heading_level = min(getattr(section, 'level', 1), 9) # Utiliser niveau de la section
                     if heading_level > 0: document.add_heading(section.title, level=heading_level)
                     # Ajouter paragraphes
                     for para in content.split('\n'):
                         if para.strip(): document.add_paragraph(para)
                     document.add_paragraph() # Espace après section
                 except Exception as e_docx: log.error(f"Error adding section '{section.title}' to DOCX: {e_docx}")
            else:
                section.status = "failed"
                log.error(f"Skipping section '{section.title}' in DOCX due to generation failure.")

            # Traiter les sous-sections
            if hasattr(section, 'subsections') and section.subsections:
                 for subsection in section.subsections:
                      process_and_add_section(subsection)

        # Traiter toutes les sections du plan
        if hasattr(report_plan, 'structure'):
             for top_section in report_plan.structure:
                  process_and_add_section(top_section)

        # TODO: Ajouter la bibliographie ici en utilisant ReferenceManager

        # Sauvegarder le document DOCX
        try:
            document.save(output_docx_path)
            log.info(f"Report draft saved successfully to {output_docx_path}")
        except Exception as e_save:
            log.error(f"Error saving DOCX file {output_docx_path}: {e_save}", exc_info=True)

        return report_plan # Retourne le plan mis à jour avec contenu et statuts


# --- AJOUT : Fonction pour assembler le rapport depuis le plan JSON ---
def assemble_report_from_plan(plan_filepath: str, output_docx_path: str):
    """
    Assembles a DOCX report from a ReportPlan JSON file that already contains
    the content for each section.
    """
    log.info(f"Loading report plan from: {plan_filepath}")
    memory_manager = MemoryManager() # Utiliser le manager pour charger proprement
    report_plan = memory_manager.load_report_plan(plan_filepath)

    if not report_plan:
        log.error(f"Failed to load report plan from {plan_filepath}. Assembly aborted.")
        raise FileNotFoundError(f"Report plan file not found or invalid: {plan_filepath}")

    log.info("Initializing DOCX document...")
    document = Document()

    # --- Style de base (optionnel, à adapter) ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # Justifier le texte

    # --- Ajout du Titre Principal ---
    report_title = getattr(report_plan, 'title', "Apprenticeship Report")
    document.add_heading(report_title, level=0)
    document.add_paragraph() # Espace après le titre

    # --- Fonction Récursive pour Ajouter les Sections ---
    def add_section_to_doc(section: ReportSection):
        title = getattr(section, 'title', 'Untitled Section')
        level = getattr(section, 'level', 1)
        content = getattr(section, 'content', None) # Récupérer le contenu sauvegardé

        log.debug(f"Adding section L{level}: '{title}'")
        try:
            # Ajouter le titre de la section
            heading_level = min(level, 9) # Limiter le niveau de titre à 9
            if heading_level > 0:
                document.add_heading(title, level=heading_level)

            # Ajouter le contenu de la section s'il existe
            if content and isinstance(content, str) and content.strip():
                # Ajouter les paragraphes
                for para_text in content.split('\n'):
                    if para_text.strip(): # Éviter les paragraphes vides
                        p = document.add_paragraph(para_text.strip())
                        p.style = document.styles['Normal'] # Appliquer le style normal
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif title.lower() not in ["bibliography", "appendices (optional)"]:
                # Ajouter un placeholder si le contenu est manquant (sauf pour biblio/annexes)
                log.warning(f"Content missing for section '{title}' (ID: {getattr(section, 'section_id', 'N/A')}). Adding placeholder.")
                p = document.add_paragraph(f"[Content for '{title}' to be added]")
                p.style = document.styles['Normal']
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                 # Placeholder spécifique pour Biblio/Annexes si vide
                 p = document.add_paragraph(f"[{title} section content goes here]")
                 p.style = document.styles['Normal']
                 p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


            document.add_paragraph() # Ajouter un espace après chaque section

        except Exception as e_docx:
            log.error(f"Error adding section '{title}' to DOCX: {e_docx}", exc_info=True)

        # Traiter les sous-sections récursivement
        if hasattr(section, 'subsections') and section.subsections:
            for subsection in section.subsections:
                add_section_to_doc(subsection)

    # --- Lancer l'assemblage depuis la racine ---
    if hasattr(report_plan, 'structure'):
        for top_section in report_plan.structure:
            add_section_to_doc(top_section)
    else:
        log.warning("Report plan structure is empty. The generated document will be minimal.")

    # --- Sauvegarder le Document Final ---
    try:
        log.info(f"Saving assembled report to: {output_docx_path}")
        document.save(output_docx_path)
        log.info("DOCX file saved successfully.")
    except Exception as e_save:
        log.error(f"Error saving final DOCX file {output_docx_path}: {e_save}", exc_info=True)
        raise # Renvoyer l'erreur pour que main.py puisse la logger

# --- FIN AJOUT ---