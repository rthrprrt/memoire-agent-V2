# document_processor.py (Version attendant le format AAAA-MM-JJ.docx)

import os
import docx # python-docx library
import config
from data_models import JournalEntry
import datetime
import logging
import hashlib
# This line imports Optional, List, and Tuple for type hinting
from typing import List, Tuple, Optional

# Configure logging
# Mise à jour du format pour inclure le nom du module pour plus de clarté
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - [%(module)s] - %(message)s')

def extract_text_from_docx(file_path: str) -> str:
    """Extracts all text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Optionnel : Extraire aussi le texte des tableaux si nécessaire
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

def parse_date_from_filename(filename: str) -> Optional[datetime.date]:
    """
    Attempts to parse a date from the filename using the format specified
    in config.DATE_FILENAME_FORMAT (expected: '%Y-%m-%d').
    """
    try:
        # Extract filename without extension
        base_name = os.path.splitext(filename)[0]
        # Attempt to parse using the defined format in config.py
        return datetime.datetime.strptime(base_name, config.DATE_FILENAME_FORMAT).date()
    except ValueError:
        # Log clearly if the format doesn't match
        logging.warning(f"Could not parse date from filename: '{filename}'. "
                        f"Expected format '{config.DATE_FILENAME_FORMAT}' (e.g., YYYY-MM-DD.docx). "
                        f"Please rename the file accordingly.")
        return None # Indicate failure to parse date
    except Exception as e:
        # Catch any other unexpected errors during parsing
        logging.error(f"Unexpected error parsing date from filename '{filename}': {e}")
        return None

def generate_entry_id(file_path: str, text: str) -> str:
    """Generates a unique ID for the journal entry."""
    # Use a hash of the file path and first part of content for robustness
    hasher = hashlib.md5()
    hasher.update(file_path.encode('utf-8'))
    hasher.update(text[:500].encode('utf-8')) # Hash first 500 chars
    return hasher.hexdigest()

def process_journal_file(file_path: str) -> Optional[JournalEntry]:
    """Processes a single journal file into a JournalEntry object."""
    logging.info(f"Processing journal file: {file_path}")
    filename = os.path.basename(file_path)
    raw_text = extract_text_from_docx(file_path)

    if not raw_text:
        logging.warning(f"No text extracted from '{filename}'. Skipping.")
        return None

    entry_date = parse_date_from_filename(filename)
    if not entry_date:
        # If date parsing fails, log the error and skip the entry
        logging.error(f"Date could not be determined for '{filename}' due to filename format mismatch. Skipping entry.")
        return None

    entry_id = generate_entry_id(file_path, raw_text)

    # Create the JournalEntry object using the Pydantic model
    entry = JournalEntry(
        entry_id=entry_id,
        date=entry_date,
        raw_text=raw_text,
        source_file=filename
        # processed_text, tags, competencies_identified, projects_mentioned
        # will be populated by later processing steps
    )
    logging.info(f"Successfully processed '{filename}' for date {entry_date}.")
    return entry

def process_all_journals(journal_dir: str = config.JOURNAL_DIR) -> List[JournalEntry]:
    """Processes all valid DOCX files (with correct filename format) in the specified directory."""
    entries: List[JournalEntry] = []
    logging.info(f"Starting journal processing in directory: {journal_dir}")
    if not os.path.isdir(journal_dir):
        logging.error(f"Journal directory not found: {journal_dir}")
        return entries # Return empty list

    for filename in os.listdir(journal_dir):
        # Check for .docx extension and ignore temporary files (like those opened in Word)
        if filename.lower().endswith(".docx") and not filename.startswith('~'):
            file_path = os.path.join(journal_dir, filename)
            if os.path.isfile(file_path): # Ensure it's actually a file
                entry = process_journal_file(file_path)
                if entry:
                    entries.append(entry)
            else:
                logging.warning(f"Found item '{filename}' which is not a file. Skipping.")

    logging.info(f"Finished processing directory. Found {len(entries)} valid journal entries.")
    # Sort entries chronologically by date
    if entries:
        entries.sort(key=lambda x: x.date)
    return entries

# --- Text Chunking ---
def chunk_text(text: str, chunk_size: int = config.CHUNK_SIZE, chunk_overlap: int = config.CHUNK_OVERLAP) -> List[str]:
    """Splits text into overlapping chunks based on character count."""
    if not text:
        return []
    if chunk_overlap < 0:
         logging.warning(f"Chunk overlap ({chunk_overlap}) is negative. Setting to 0.")
         chunk_overlap = 0
    if chunk_overlap >= chunk_size:
        logging.warning(f"Chunk overlap ({chunk_overlap}) >= chunk size ({chunk_size}). Setting overlap to {chunk_size // 4}.")
        chunk_overlap = chunk_size // 4 # Prevent infinite loops or nonsensical overlap

    chunks = []
    start_index = 0
    text_len = len(text)

    while start_index < text_len:
        end_index = min(start_index + chunk_size, text_len)
        chunks.append(text[start_index:end_index])

        # Move start_index for the next chunk
        next_start_index = start_index + chunk_size - chunk_overlap

        # Ensure forward progress to avoid infinite loops
        if next_start_index <= start_index:
            if start_index >= text_len - 1: # Break if already at the end
                break
            start_index += 1 # Force minimal progress
        else:
            start_index = next_start_index

    # Remove potentially empty chunks resulting from splitting logic
    return [chunk for chunk in chunks if chunk.strip()]